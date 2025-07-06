import logging
from docx import Document
from docx.oxml.ns import qn
from typing import List, Dict, Any


def load_docx(file_path: str) -> Dict[str, Any]:
    try:
        doc = Document(file_path)
        result = {
            'content': '',
            'metadata': {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'num_pages': len(doc.paragraphs),
                'num_tables': len(doc.tables)
            },
            'sections': []
        }

        # 提取文本段落
        text_content = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_content.append(f"[段落{i + 1}] {para.text}")

        # 提取表格
        for i, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_content.append("\t".join(row_data))
            text_content.append(f"[表格{i + 1}]\n" + "\n".join(table_content))

        result['content'] = "\n\n".join(text_content)
        result['size'] = len(result['content'])
        result['status'] = 'success'

        return result

    except Exception as e:
        logging.error(f"Error loading DOCX {file_path}: {str(e)}")
        return {
            'content': f"Error: {str(e)}",
            'metadata': {},
            'sections': [],
            'size': 0,
            'status': 'error'
        }
